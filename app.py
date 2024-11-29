from flask import Flask, render_template, request, redirect, url_for, flash, session
from pymongo import MongoClient
import os
import uuid
import hashlib
from datetime import datetime
import shutil
from werkzeug.utils import secure_filename
from PyPDF2 import PdfFileReader
from docx import Document
import pandas as pd
from PIL import Image
import piexif
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image, ExifTags
import re
import openpyxl
from PIL import Image, ExifTags
import io
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle


app = Flask(__name__)
app.secret_key = "supersecretkey"

# MongoDB Configuration
mongo_uri = "mongodb://127.0.0.1:27017/"
client = MongoClient(mongo_uri)
db = client['Lakshmeesh']
tenant_collection = db['tenant']

LOCAL_STORAGE_BASE = 'local_storage'


def get_user_module_path(tenant_id, username, module_name, workspace_name=None):
    """Construct the local storage path based on tenant, user, module, and optional workspace."""
    path = os.path.join(LOCAL_STORAGE_BASE, tenant_id, username, module_name)
    if workspace_name:
        path = os.path.join(path, workspace_name)
    return path


def create_module_collection(tenant_id, username, module_name):
    """Create a module-specific collection for a user."""
    collection_name = f"{tenant_id}_{username}_{module_name}"
    if collection_name not in db.list_collection_names():
        db.create_collection(collection_name)
    return db[collection_name]


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        tenant_id = request.form.get('tenantId')
        username = request.form.get('username')
        password = request.form.get('password')

        if tenant_id and username and password:
            user_data = tenant_collection.find_one({
                "tenantId": tenant_id,
                "users": {
                    "$elemMatch": {
                        "username": username,
                        "password": password
                    }
                }
            })
            if user_data:
                session['tenantId'] = tenant_id
                session['username'] = username
                flash("Logged in successfully!", "success")
                return redirect(url_for('select_folder'))
            else:
                flash("Login failed. Please check your credentials.", "error")
    return render_template('index.html')


@app.route('/select_folder', methods=['GET', 'POST'])
def select_folder():
    tenant_id = session.get('tenantId')
    username = session.get('username')

    if not tenant_id or not username:
        flash("Session expired, please log in again.", "error")
        return redirect(url_for('index'))

    modules = ["endoc", "ensight", "envision"]

    if request.method == 'POST':
        module_name = request.form.get('folder')
        if module_name:
            session['module_name'] = module_name
            return redirect(url_for('workspace'))
        else:
            flash("Please select a module.", "error")

    return render_template('select_folder.html', modules=modules)


@app.route('/workspace', methods=['GET', 'POST'])
def workspace():
    tenant_id = session.get('tenantId')
    username = session.get('username')
    module_name = session.get('module_name')

    if not all([tenant_id, username, module_name]):
        flash("Session expired or module not selected.", "error")
        return redirect(url_for('select_folder'))

    if request.method == 'POST':
        workspace_name = request.form.get('workspace_name')
        selected_workspace = request.form.get('selected_workspace')

        if workspace_name:
            # Check if the workspace already exists
            existing_workspace = tenant_collection.find_one({
                "tenantId": tenant_id,
                "module": module_name,
                "workspaceName": workspace_name,
                "username": username
            })

            if existing_workspace:
                flash("Workspace already exists.", "error")
            else:
                # Assign UUIDs for workspace and user
                workspace_uuid = str(uuid.uuid4())

                # Insert workspace metadata
                tenant_collection.insert_one({
                    "tenantId": tenant_id,
                    "module": module_name,
                    "workspaceName": workspace_name,
                    "workspaceUUID": workspace_uuid,
                    "username": username,
                    "files": []
                })

                # Create the workspace directory in local storage
                workspace_path = get_user_module_path(tenant_id, username, module_name, workspace_name)
                os.makedirs(workspace_path, exist_ok=True)

                flash("Workspace created successfully!", "success")
                session['workspace_name'] = workspace_name
                return redirect(url_for('upload_files'))

        elif selected_workspace:
            # Handle existing workspace selection
            session['workspace_name'] = selected_workspace
            flash(f"Workspace '{selected_workspace}' selected.", "success")
            return redirect(url_for('upload_files'))

        else:
            flash("Please enter a workspace name or select an existing workspace.", "error")

    # Fetch all workspaces created by the user for the selected module
    workspaces = tenant_collection.find({
        "tenantId": tenant_id,
        "module": module_name,
        "username": username
    })
    return render_template('create_workspace_in_folder.html', workspaces=workspaces)

@app.route('/delete_workspace/<workspace_uuid>', methods=['POST'])
def delete_workspace(workspace_uuid):
    tenant_id = session.get('tenantId')
    user_name = session.get('username')
    module_name = session.get('module_name')  # Get the module_name from the session

    if not tenant_id or not user_name or not module_name:  # Check if module_name is also available
        flash("Session expired, please log in again.", "error")
        return redirect(url_for('index'))

    # Check if workspace exists for this user by workspaceUUID and workspaceName
    user_data = tenant_collection.find_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "username": user_name,
            "workspaceUUID": workspace_uuid  # Ensure we're targeting the correct workspace
        }
    )

    if not user_data:
        flash('Workspace does not exist or you do not have access to it.', 'error')
        return redirect(url_for('workspace'))

    # Since workspaceUUID is a direct field, you can directly access it.
    workspace_name = user_data.get('workspaceName')

    if not workspace_name:
        flash('Workspace name not found.', 'error')
        return redirect(url_for('workspace'))

    # Construct local directory path
    workspace_dir = get_user_module_path(tenant_id, user_name, module_name, workspace_name)

    # Remove local directory if it exists
    if os.path.exists(workspace_dir):
        try:
            shutil.rmtree(workspace_dir)
            flash('Workspace and its files deleted successfully.', 'success')
        except Exception as e:
            flash(f'Failed to delete local workspace directory: {str(e)}', 'error')

    # Now, delete the entire workspace document from MongoDB
    result = tenant_collection.delete_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "username": user_name,
            "workspaceUUID": workspace_uuid  # This will match the specific workspace document
        }
    )

    if result.deleted_count > 0:
        flash('Workspace metadata deleted successfully from MongoDB.', 'success')
    else:
        flash('Failed to delete workspace metadata from MongoDB.', 'error')

    return redirect(url_for('workspace'))

def extract_metadata_for_endoc(file):
    metadata = {}

    
    url_pattern = r'https?://[^\s]+'

    if file.filename.endswith(".pdf"):
        try:
            
            file.seek(0)  
            reader = PdfReader(file)
            document_info = reader.metadata  

            metadata['title'] = document_info.get('/Title', 'Unknown') if document_info else 'Unknown'
            metadata['author'] = document_info.get('/Author', 'Unknown') if document_info else 'Unknown'
            metadata['num_pages'] = len(reader.pages)

            
            all_urls = []
            for page in reader.pages:
                text = page.extract_text() or ""
                urls = re.findall(url_pattern, text)
                all_urls.extend(urls)

            metadata['urls'] = list(set(all_urls))  # Remove duplicates and store in metadata

        except Exception as e:
            metadata['error'] = f"Error extracting PDF metadata: {str(e)}"

    elif file.filename.endswith((".doc", ".docx")):
        try:
            # Use python-docx for Word documents
            file.seek(0)  # Reset file pointer
            doc = Document(file)

            metadata['title'] = doc.core_properties.title or "Unknown"
            metadata['author'] = doc.core_properties.author or "Unknown"
            metadata['num_paragraphs'] = len(doc.paragraphs)

            # Extract and find URLs from the Word document
            all_urls = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                urls = re.findall(url_pattern, text)
                all_urls.extend(urls)

            metadata['urls'] = list(set(all_urls))  # Remove duplicates and store in metadata

        except Exception as e:
            metadata['error'] = f"Error extracting DOC/DOCX metadata: {str(e)}"

    return metadata


def extract_metadata_for_ensight(file):
    metadata = {}

    # Regex for URL detection
    url_pattern = r'https?://[^\s]+'

    if file.filename.endswith((".xlsx", ".xlsm", ".csv")):
        try:
            if file.filename.endswith((".xlsx", ".xlsm")):
                # Read Excel file with all sheets
                df = pd.read_excel(file, sheet_name=None)
            else:
                # Read CSV file as a single sheet
                df = {"Sheet1": pd.read_csv(file)}

            for sheet_name, sheet_data in df.items():
                sheet_metadata = {"columns": list(sheet_data.columns), "urls": [], "hyperlinks": {}}

                # Extract URLs from cell values
                for column in sheet_data.columns:
                    column_values = sheet_data[column].astype(str)
                    for value in column_values:
                        urls = re.findall(url_pattern, value)
                        sheet_metadata["urls"].extend(urls)  # Collect URLs

                # Check for hyperlinks (specific to Excel files)
                if file.filename.endswith((".xlsx", ".xlsm")):
                    # Reload Excel as openpyxl to access hyperlinks
                    workbook = openpyxl.load_workbook(file, data_only=True)
                    if sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        for row in sheet.iter_rows():
                            for cell in row:
                                if cell.hyperlink:
                                    sheet_metadata["hyperlinks"][cell.coordinate] = cell.hyperlink.target

                # Remove duplicate URLs
                sheet_metadata["urls"] = list(set(sheet_metadata["urls"]))
                metadata[sheet_name] = sheet_metadata

        except Exception as e:
            metadata['error'] = f"Error extracting Excel/CSV metadata: {str(e)}"

    return metadata



from PIL import Image, ExifTags
import io

def extract_metadata_for_envision(file):
    metadata = {}

    if file.filename.endswith((".jpeg", ".png", ".jpg", ".bmp", ".tiff", ".webp")):  # Added more formats
        try:
            file.seek(0)  # Reset file pointer
            img = Image.open(file)

            # Basic image metadata
            metadata['width'], metadata['height'] = img.size
            metadata['format'] = img.format
            metadata['mode'] = img.mode  # Image mode (e.g., RGB, L)
            metadata['size_in_kb'] = len(file.read()) // 1024  # Approximate size in KB


            # Dominant color (if mode is RGB)
            if img.mode == "RGB":
                img.thumbnail((50, 50))  # Resize for efficiency
                pixels = list(img.getdata())
                dominant_color = max(set(pixels), key=pixels.count)
                metadata['dominant_color'] = dominant_color

            # Histogram (useful for analyzing image content)
            metadata['histogram'] = img.histogram()

        except Exception as e:
            metadata['error'] = f"Error extracting image metadata: {str(e)}"

    return metadata


import pickle

# Set up the Google Drive API credentials
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']


@app.route('/connect_google_drive', methods=['GET'])
def connect_google_drive():
    credentials = None
    creds_path = os.path.join(os.getcwd(), 'credentials.json')
    token_path = os.path.join(os.getcwd(), 'token.pickle')

    # Load existing token or generate a new one
    if os.path.exists(token_path):
        with open(token_path, 'rb') as token_file:
            credentials = pickle.load(token_file)

    if not credentials or not credentials.valid:
        if credentials and credentials.expired and credentials.refresh_token:
            credentials.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(creds_path, SCOPES)
            credentials = flow.run_local_server(port=8082)
        
        # Save credentials for future use
        with open(token_path, 'wb') as token_file:
            pickle.dump(credentials, token_file)

    # Connect to Google Drive
    drive_service = build('drive', 'v3', credentials=credentials)

    # List files in the user's Drive
    results = drive_service.files().list(pageSize=10, fields="files(id, name)").execute()
    items = results.get('files', [])

    # Debugging: Print the raw API response
    print(results)  # This will print the entire response from Google Drive API

    if not items:
        files_list = "No files found in Google Drive."
    else:
        files_list = [(file['name'], file['id']) for file in items]

        for file in items:
            try:
                # Safely extract 'name' and 'id' from the file
                name = file.get('name', 'Unnamed File')
                file_id = file.get('id', 'No ID')
                files_list.append((name, file_id))  # Appending tuple (name, file_id)
            except ValueError as e:
                print(f"Error unpacking file data: {e}")
                continue

    # Debugging step: Print files_list to check the structure
    print(f"files_list: {files_list}")

    # Ensure that the files_list is structured correctly for the template
    if isinstance(files_list, str):
        return files_list  # Return the message directly if no files are found
    elif len(files_list) > 0 and isinstance(files_list[0], tuple):
        return render_template('google_drive.html', files=files_list)
    else:
        return "Error: The files list is not structured correctly."



@app.route('/upload', methods=['GET', 'POST'])
def upload_files():
    tenant_id = session.get('tenantId')
    username = session.get('username')
    module_name = session.get('module_name')
    workspace_name = session.get('workspace_name')

    # If 'workspace_name' is passed as a query parameter, update the session
    if request.args.get('workspace_name'):
        workspace_name = request.args.get('workspace_name')
        session['workspace_name'] = workspace_name

    # Validate required fields in session
    if not all([tenant_id, username, module_name, workspace_name]):
        flash("Error: Missing required fields.", "error")
        return redirect(url_for('workspace'))

    upload_folder = get_user_module_path(tenant_id, username, module_name, workspace_name)
    os.makedirs(upload_folder, exist_ok=True)

    allowed_file_types = {
        "endoc": [".pdf", ".doc", ".docx"],
        "ensight": [".xlsx", ".csv", ".xlsm"],
        "envision": [".jpeg", ".png", ".jpg", ".bmp", ".tiff", ".webp"]
    }

    if request.method == 'POST':
        files = request.files.getlist('files')
        uploaded_files_metadata = []
        error_message = ""

        for file in files:
            file_extension = os.path.splitext(file.filename)[1].lower()

            if file_extension not in allowed_file_types.get(module_name, []):
                error_message = f"Only {', '.join(allowed_file_types[module_name])} files are allowed for the {module_name} module."
                break

            file_hash = hashlib.sha256(file.read()).hexdigest()
            file.seek(0)

            existing_file = tenant_collection.find_one(
                {
                    "tenantId": tenant_id,
                    "module": module_name,
                    "workspaceName": workspace_name,
                    "files.hash": file_hash
                }
            )
            if existing_file:
                flash(f"Duplicate file skipped: {file.filename}", "warning")
                continue

            file_path = os.path.join(upload_folder, file.filename)
            file.save(file_path)

            # Extract metadata separately based on the module
            metadata = {}
            if module_name == "endoc":
                metadata = extract_metadata_for_endoc(file)
            elif module_name == "ensight":
                metadata = extract_metadata_for_ensight(file)
            elif module_name == "envision":
                metadata = extract_metadata_for_envision(file)

            # File metadata
            file_metadata = {
                "filename": file.filename,
                "hash": file_hash,
                "upload_time": datetime.utcnow(),
                "file_path": file_path,
                "metadata": metadata  # Add extracted metadata here
            }

            uploaded_files_metadata.append(file_metadata)

        if error_message:
            flash(error_message, "error")
        elif uploaded_files_metadata:
            tenant_collection.update_one(
                {
                    "tenantId": tenant_id,
                    "module": module_name,
                    "workspaceName": workspace_name
                },
                {"$push": {"files": {"$each": uploaded_files_metadata}}}
            )
            flash("Files uploaded successfully!", "success")

    # Fetch the list of files in the correct workspace and module
    workspace_data = tenant_collection.find_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "workspaceName": workspace_name  # Ensure we query based on the selected workspace and module
        }
    )

    # Fetch only the files associated with the workspace and module
    files_in_workspace = workspace_data.get('files', []) if workspace_data else []

    return render_template('upload.html', selected_workspace=workspace_name, files_in_workspace=files_in_workspace)

@app.route('/upload_google_drive_files', methods=['POST'])
def upload_google_drive_files():
    # Get selected files from the form
    selected_file_ids = request.form.getlist('selected_files')

    if not selected_file_ids:
        flash("No files selected for upload.", "error")
        return redirect(url_for('connect_google_drive'))

    # Set up Google Drive API credentials
    creds_path = os.path.join(os.getcwd(), 'credentials.json')
    token_path = os.path.join(os.getcwd(), 'token.pickle')

    with open(token_path, 'rb') as token_file:
        credentials = pickle.load(token_file)

    drive_service = build('drive', 'v3', credentials=credentials)

    # Get tenant, user, module, and workspace details
    tenant_id = session.get('tenantId')
    username = session.get('username')
    module_name = session.get('module_name')
    workspace_name = session.get('workspace_name')

    if not all([tenant_id, username, module_name, workspace_name]):
        flash("Missing session information.", "error")
        return redirect(url_for('upload_files'))

    upload_folder = get_user_module_path(tenant_id, username, module_name, workspace_name)
    os.makedirs(upload_folder, exist_ok=True)

    uploaded_files_metadata = []

    for file_id in selected_file_ids:
        # Get file metadata from Google Drive
        file_metadata = drive_service.files().get(fileId=file_id).execute()
        file_name = file_metadata.get('name', 'Unnamed File')

        # Download file content
        request = drive_service.files().get_media(fileId=file_id)
        file_path = os.path.join(upload_folder, file_name)

        with open(file_path, 'wb') as f:
            downloader = MediaIoBaseDownload(f, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()

        # Calculate file hash
        with open(file_path, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()

        # Check for duplicate files in DB
        existing_file = tenant_collection.find_one({
            "tenantId": tenant_id,
            "module": module_name,
            "workspaceName": workspace_name,
            "files.hash": file_hash
        })
        if existing_file:
            flash(f"Duplicate file skipped: {file_name}", "warning")
            continue

        # Add file metadata
        uploaded_files_metadata.append({
            "filename": file_name,
            "hash": file_hash,
            "upload_time": datetime.utcnow(),
            "file_path": file_path
        })

    # Update DB with uploaded files
    if uploaded_files_metadata:
        tenant_collection.update_one(
            {"tenantId": tenant_id, "module": module_name, "workspaceName": workspace_name},
            {"$push": {"files": {"$each": uploaded_files_metadata}}}
        )
        flash("Files uploaded successfully!", "success")
    else:
        flash("No new files were uploaded.", "info")

    return redirect(url_for('upload_files'))


@app.route('/update_file/<file_hash>', methods=['POST'])
def update_file(file_hash):
    tenant_id = session.get('tenantId')
    username = session.get('username')
    module_name = session.get('module_name')
    workspace_name = session.get('workspace_name')

    if not all([tenant_id, username, module_name, workspace_name]):
        flash("Error: Missing required fields.", "error")
        return redirect(url_for('upload_files'))

    upload_folder = get_user_module_path(tenant_id, username, module_name, workspace_name)

    # Find the file in the database
    workspace_data = tenant_collection.find_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "workspaceName": workspace_name,
            "files.hash": file_hash
        }
    )
    if not workspace_data:
        flash("Error: File not found for update.", "error")
        return redirect(url_for('upload_files'))

    # Get the uploaded file
    new_file = request.files.get('file')
    if not new_file:
        flash("Error: No file provided for update.", "error")
        return redirect(url_for('upload_files'))

    # Validate file type
    file_extension = os.path.splitext(new_file.filename)[1].lower()
    allowed_file_types = {
        "endoc": [".pdf", ".doc", ".docx"],
        "ensight": [".xlsx", ".csv", ".xlsm"],
        "envision": [".jpeg", ".png", ".jpg", ".bmp", ".tiff", ".webp"]
    }
    if file_extension not in allowed_file_types.get(module_name, []):
        flash(f"Error: Only {', '.join(allowed_file_types[module_name])} files are allowed for the {module_name} module.", "error")
        return redirect(url_for('upload_files'))

    # Generate hash and check for duplicates
    new_file_hash = hashlib.sha256(new_file.read()).hexdigest()
    new_file.seek(0)
    if tenant_collection.find_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "workspaceName": workspace_name,
            "files.hash": new_file_hash
        }
    ):
        flash("Error: Duplicate file detected. Update aborted.", "error")
        return redirect(url_for('upload_files'))

    # Save the new file
    new_file_path = os.path.join(upload_folder, new_file.filename)
    new_file.save(new_file_path)

    # Extract metadata for the new file
    metadata = {}
    if module_name == "endoc":
        metadata = extract_metadata_for_endoc(new_file)
    elif module_name == "ensight":
        metadata = extract_metadata_for_ensight(new_file)
    elif module_name == "envision":
        metadata = extract_metadata_for_envision(new_file)

    # Update the database: Replace the file's entry
    tenant_collection.update_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "workspaceName": workspace_name,
            "files.hash": file_hash
        },
        {
            "$set": {
                "files.$.filename": new_file.filename,
                "files.$.hash": new_file_hash,
                "files.$.upload_time": datetime.utcnow(),
                "files.$.file_path": new_file_path,
                "files.$.metadata": metadata
            }
        }
    )

    # Delete the old file from local storage
    old_file_path = os.path.join(upload_folder, workspace_data['files'][0]['filename'])
    if os.path.exists(old_file_path):
        os.remove(old_file_path)

    flash("File updated successfully!", "success")
    return redirect(url_for('upload_files'))




@app.route('/delete_file/<file_hash>', methods=['POST'])
def delete_file(file_hash):
    tenant_id = session.get('tenantId')
    username = session.get('username')
    module_name = session.get('module_name')
    workspace_name = session.get('workspace_name')

    if not all([tenant_id, username, module_name, workspace_name]):
        flash("Error: Missing required fields.", "error")
        return redirect(url_for('workspace'))

    # Find the file metadata in the database
    workspace_data = tenant_collection.find_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "username": username,
            "workspaceName": workspace_name
        }
    )
    file_metadata = next((file for file in workspace_data.get('files', []) if file['hash'] == file_hash), None)

    if not file_metadata:
        flash("File not found.", "error")
        return redirect(url_for('upload_files'))

    # Delete the file from local storage
    file_path = os.path.join(get_user_module_path(tenant_id, username, module_name, workspace_name), file_metadata['filename'])
    if os.path.exists(file_path):
        os.remove(file_path)

    # Remove the file metadata from the MongoDB collection
    tenant_collection.update_one(
        {
            "tenantId": tenant_id,
            "module": module_name,
            "username": username,
            "workspaceName": workspace_name
        },
        {"$pull": {"files": {"hash": file_hash}}}
    )

    flash(f"File '{file_metadata['filename']}' deleted successfully.", "success")
    return redirect(url_for('upload_files'))



# @app.route('/delete_workspace/<workspace_uuid>', methods=['POST'])
# def delete_workspace(workspace_uuid):
#     tenant_id = session.get('tenantId')
#     user_name = session.get('username')
#     module_name = session.get('module_name')  # Get the module_name from the session

#     if not tenant_id or not user_name or not module_name:  # Check if module_name is also available
#         flash("Session expired, please log in again.", "error")
#         return redirect(url_for('index'))

#     # Check if workspace exists for this user by workspaceUUID and workspaceName
#     user_data = tenant_collection.find_one(
#         {
#             "tenantId": tenant_id,
#             "module": module_name,
#             "username": user_name,
#             "workspaceUUID": workspace_uuid  # Ensure we're targeting the correct workspace
#         }
#     )

#     if not user_data:
#         flash('Workspace does not exist or you do not have access to it.', 'error')
#         return redirect(url_for('workspace'))

#     # Since workspaceUUID is a direct field, you can directly access it.
#     workspace_name = user_data.get('workspaceName')

#     if not workspace_name:
#         flash('Workspace name not found.', 'error')
#         return redirect(url_for('workspace'))

#     # Construct local directory path
#     workspace_dir = get_user_module_path(tenant_id, user_name, module_name, workspace_name)

#     # Remove local directory if it exists
#     if os.path.exists(workspace_dir):
#         try:
#             shutil.rmtree(workspace_dir)
#             flash('Workspace and its files deleted successfully.', 'success')
#         except Exception as e:
#             flash(f'Failed to delete local workspace directory: {str(e)}', 'error')

#     # Now, delete the entire workspace document from MongoDB
#     result = tenant_collection.delete_one(
#         {
#             "tenantId": tenant_id,
#             "module": module_name,
#             "username": user_name,
#             "workspaceUUID": workspace_uuid  # This will match the specific workspace document
#         }
#     )

#     if result.deleted_count > 0:
#         flash('Workspace metadata deleted successfully from MongoDB.', 'success')
#     else:
#         flash('Failed to delete workspace metadata from MongoDB.', 'error')

#     return redirect(url_for('workspace'))


@app.route('/upload_result')
def upload_result():
    uploaded = request.args.getlist('uploaded')
    duplicates = request.args.getlist('duplicates')

    uploaded_files = ', '.join(uploaded) if uploaded else "None"
    duplicate_files = ', '.join(duplicates) if duplicates else "None"

    message = (
        "<span style='color: red;'>"
        "The file you're trying to upload already exists in the workspace. "
        "If you want to replace it with a new one, please use the update option."
        "</span>"
    )

    return f"""
        <p>Files uploaded successfully: {uploaded_files}</p><br><br><br>
        <p>Duplicate files: {duplicate_files}</p>
        <p>{message}</p>
    """



@app.route('/upload_duplicates')
def upload_duplicates():
    return "The file your trying to upload already exist in the workspace , if you want to replace it with a new one please use the update option."






@app.route('/upload_success')
def upload_success():
    return "Files uploaded successfully!"


if __name__ == '__main__':
    os.makedirs(LOCAL_STORAGE_BASE, exist_ok=True)
    app.run(debug=True)
